import json
from pickle import FALSE
from docx import Document
from docx.shared import RGBColor

# for testing
import csv

# dict av enkeltbesvarelse -> void + rapport i docx-format
def generate_report_t1v20(data, codebook_path, outputPath, respondentID):
    
    # loads codebook as dict
    with open(codebook_path, encoding="utf-8") as f:
        codebook = json.load(f)

    # creates new document and adds heading
    document = Document()
    document.add_heading('Pasientrapportert kartlegging ved oppstart i arbeidsrettet rehabilitering', 1)
    document.add_heading('Avdeling for fysikalsk medisin og forebygging, Sørlandet sykehus HF', 3)

    ### INNLEDNING ###    
    write_intro_t1(data, codebook, document, respondentID)

    ### OPPSUMMERING ###
    write_summary_t1v20(data, codebook, document)
    
    ### SMERTE ###
    write_pain_variables(data, codebook, document)
    write_earlier_treatment(data, codebook, document)

    ### ARBEID OG YTELSER ###
    write_work_related_t1v20(data, codebook, document)

    ### Livskvalitet
    write_quality_of_life(data, codebook, document)
    
    ### FYSISK AKTIVITET, HØYDE, VEKT, KOSTHOLD ###
    write_exercise_and_more(data, codebook, document)

    ### SCL-10 ###
    write_scl(data, codebook, document)

    ### ISI ###
    write_isi(data, codebook, document)


    # saves document to file
    document.save(outputPath + str(respondentID) + ".docx")
    print("T1_v2.0-rapport generert for " + str(respondentID) + " (" + data["fnr"][0:6] + " " + data["fnr"][6:] + ")")

    if(data["tilbakemeldinger"] != ""):
        print("Respondent har gitt tilbakemelding: " + data["tilbakemeldinger"])



def generate_report_t2v10(data, codebook_path, outputPath, respondentID):
        # loads codebook as dict
    with open(codebook_path, encoding="utf-8") as f:
        codebook = json.load(f)

    # creates new document and adds heading
    document = Document()
    document.add_heading('Pasientrapportert kartlegging ved avslutning i arbeidsrettet rehabilitering', 1)
    document.add_heading('Avdeling for fysikalsk medisin og forebygging, Sørlandet sykehus HF', 3)

    ### INNLEDNING ###    
    write_intro_t2(data, codebook, document, respondentID)

    ### OPPSUMMERING ###
    write_summary_t2v10(data, codebook, document)
    
    ### SMERTE ###
    write_pain_variables(data, codebook, document)
    
    ### ARBEID OG YTELSER ###
    write_work_related_t2v10(data, codebook, document)

    ### Livskvalitet
    write_quality_of_life(data, codebook, document)
    
    ### FYSISK AKTIVITET, HØYDE, VEKT, KOSTHOLD ###
    write_exercise_and_more(data, codebook, document)

    ### SCL-10 ###
    write_scl(data, codebook, document)

    ### ISI ###
    write_isi(data, codebook, document)

    # saves document to file
    document.save(outputPath + str(respondentID) + ".docx")
    print("T2_v1.0-rapport generert for " + str(respondentID) + " (" + data["fnr"][0:6] + " " + data["fnr"][6:] + ")")

    if(data["tilbakemeldinger"] != ""):
        print("Respondent har gitt tilbakemelding: " + data["tilbakemeldinger"])


####################
### INTRODUKSJON ###
####################
#region Intro

def write_intro_t1(data, codebook, document, respondentID):
    # Fnr
    p = document.add_paragraph("Respondent-ID: " + str(respondentID))
    
    # Dato for utfylling
    p.add_run("\n")
    p.add_run("Dato utfylt: " + data["Opprettet"])
    
    document.add_heading("Demografi", 2)
    demografi = document.add_paragraph("")
    # Morsmål og lese-/skrivevansker
    if (data["morsmaal"] == "annet"):
        demografi.add_run("Morsmål: " + data["morsmaal-tekst"].capitalize())
        write_var_snippet_and_response("sprakvansker", data, codebook, demografi)
        demografi.add_run("\n")

    # Sivilstatus, barn, husholdning
    write_var_snippet_and_response("sivilstatus", data, codebook, demografi, True, False)
    demografi.add_run("\nBarn: " + data["barn"])
    demografi.add_run("\nPersoner i husholdningen (i tillegg til pas): " + data["antall-i-husholdning"])

def write_intro_t2(data, codebook, document, respondentID):
    # Fnr
    p = document.add_paragraph("Respondent-ID: " + str(respondentID))
    
    # Dato for utfylling
    p.add_run("\n")
    p.add_run("Dato utfylt: " + data["Opprettet"])   
 

#endregion

####################
### OPPSUMMERING ###
####################
#region Oppsummering

def write_summary_t1v20(data, codebook, document):
    # Diverse
    document.add_heading('Oppsummering', 2)
    oppsummering = document.add_paragraph("Viktigste problem: " + data["viktigste-problem"])
    oppsummering.add_run("\nPasientens tanker om årsak til plagene: " + data["aarsak"])
    oppsummering.add_run("\nOppfølging pasienten tror vil være mest nyttig: " + data["type-hjelp"])
    oppsummering.add_run("\nEndringer i jobbsituasjon pasienten har tro på: " + data["endringer-jobbsit-rtw"])
       

    # Erstatningssak og uførsøknad
    if data["erstatningssak"] == "ja":
        oppsummering.add_run("\nPågående erstatningssak")
    else:
        oppsummering.add_run("\nIngen pågående erstatningssak")

    if data["sokt-ufor"] == "ja":
        oppsummering.add_run("\nHar søkt eller planlegger å søke uførpensjon")
    else:
        oppsummering.add_run("\nVurderer ikke å søke ufør")

    # Jobbstatus
    #yrke = codebook["yrke"]["responses"][data["yrke"]] + " - " + data["yrke-fritekst"]
    if(data["arbeidsforhold"] == "ja"):
        oppsummering.add_run("\nNåværende yrke: " + data["yrke-fritekst"])
    else:
        oppsummering.add_run("\nSiste hovedyrke (har ikke arbeidsforhold): " + data["yrke-fritekst"])
    
    jobbstatus = ""
    if data["sm-aap"] == "nei": jobbstatus = "ikke sykemeldt"
    elif data["sm-aap"] == "delvis-sm": jobbstatus = "delvis sykemeldt"
    elif data["sm-aap"] == "fullt-sm": jobbstatus = "fullt sykemeldt"
    elif data["sm-aap"] == "aap": jobbstatus = "AAP"
    oppsummering.add_run(" (" + jobbstatus + ")")


    # WPI og SSS
    oppsummering.add_run("\nWPI (0-19): " + str((wpi_score(data))))
    oppsummering.add_run("\nSSS (0-12): " + str((sss_score(data, codebook))))
    oppsummering.add_run("\nSum SSS + WPI (0-31): " + str((sss_score(data, codebook) + wpi_score((data)))))


    # Antall smerteregioner
    oppsummering.add_run("\nAntall smerteregioner: " + str(number_of_pain_regions(data)))
    

    # Fibro = (WPI >=7 & SSS >=5 || WPI >=4 & SSS >=9) & >=4 kroppsregioner & >=3 mnd
    if(oppfyller_fibrokriterier(data, codebook)):
        oppsummering.add_run("\nPositivt svar på samlede kriterier for utbredte smerter")
    else:
        oppsummering.add_run("\nNegativt svar på samlede kriterier for utbredte smerter")


    # SCL-10
    oppsummering.add_run("\nSCL-10 (fra 1,0 = laveste skåre, til 4,0 = høyeste skåre): " + str(scl_score(data)))
    if (scl_score(data) > 1.85):
        oppsummering.add_run(" (indikerer psykiske plager)")
    else:
        oppsummering.add_run(" (indikerer fravær av psykiske plager)")


    # ISI
    oppsummering.add_run("\nISI: ")
    isi = isi_score(data)
    isi_normalized = isi / 28
    if(isi < 8):
        oppsummering.add_run(str(isi) + " (indikerer fravær av insomni)").font.color.rgb = RGBColor(51, 102, 0)
    elif (isi <15):
        oppsummering.add_run(str(isi) + " (indikerer subklinisk insomni)").font.color.rgb = RGBColor(153, 102, 51)
    elif (isi <22):
        oppsummering.add_run(str(isi) + " (indikerer moderat insomni)").font.color.rgb = RGBColor(204, 102, 0)
    else:
        oppsummering.add_run(str(isi) + " (indikerer alvorlig insomni)").font.color.rgb = RGBColor(204, 0, 0)


        
def write_summary_t2v10(data, codebook, document):
    # Diverse
    document.add_heading('Oppsummering', 2)
    oppsummering = document.add_paragraph("")
    oppsummering.add_run("Pasientens tanker om årsak til plagene: " + data["aarsak"])
       

    # Uførsøknad
    if data["sokt-ufor"] == "ja":
        oppsummering.add_run("\nHar søkt eller planlegger å søke uførpensjon")
    else:
        oppsummering.add_run("\nVurderer ikke å søke ufør")


    # Jobbstatus
    if(data["arbeidsforhold"] == "ja"):
        oppsummering.add_run("\nHar arbeidsforhold ")
    else:
        oppsummering.add_run("\nHar ikke arbeidsforhold ")
    jobbstatus = ""
    if data["sm-aap"] == "nei": jobbstatus = "ikke sykemeldt"
    elif data["sm-aap"] == "delvis-sm": jobbstatus = "delvis sykemeldt"
    elif data["sm-aap"] == "fullt-sm": jobbstatus = "fullt sykemeldt"
    elif data["sm-aap"] == "aap": jobbstatus = "AAP"
    oppsummering.add_run(" (" + jobbstatus + ")")
    

    # WPI og SSS
    oppsummering.add_run("\nWPI (0-19): " + str((wpi_score(data))))
    oppsummering.add_run("\nSSS (0-12): " + str((sss_score(data, codebook))))
    oppsummering.add_run("\nSum SSS + WPI (0-31): " + str((sss_score(data, codebook) + wpi_score((data)))))


    # Antall smerteregioner
    oppsummering.add_run("\nAntall smerteregioner: " + str(number_of_pain_regions(data)))
    

    # Fibro = (WPI >=7 & SSS >=5 || WPI >=4 & SSS >=9) & >=4 kroppsregioner & >=3 mnd
    if(oppfyller_fibrokriterier(data, codebook)):
        oppsummering.add_run("\nPositivt svar på samlede kriterier for utbredte smerter")
    else:
        oppsummering.add_run("\nNegativt svar på samlede kriterier for utbredte smerter")


    # SCL-10
    oppsummering.add_run("\nSCL-10 (fra 1,0 = laveste skåre, til 4,0 = høyeste skåre): " + str(scl_score(data)))
    if (scl_score(data) > 1.85):
        oppsummering.add_run(" (indikerer psykiske plager)")
    else:
        oppsummering.add_run(" (indikerer fravær av psykiske plager)")


    # ISI
    oppsummering.add_run("\nISI: ")
    isi = isi_score(data)
    isi_normalized = isi / 28
    if(isi < 8):
        oppsummering.add_run(str(isi) + " (indikerer fravær av insomni)").font.color.rgb = RGBColor(51, 102, 0)
    elif (isi <15):
        oppsummering.add_run(str(isi) + " (indikerer subklinisk insomni)").font.color.rgb = RGBColor(153, 102, 51)
    elif (isi <22):
        oppsummering.add_run(str(isi) + " (indikerer moderat insomni)").font.color.rgb = RGBColor(204, 102, 0)
    else:
        oppsummering.add_run(str(isi) + " (indikerer alvorlig insomni)").font.color.rgb = RGBColor(204, 0, 0)


#endregion

######################
### PAIN VARIABELS ###
######################
#region Smerte

def write_pain_variables(data, codebook, document):
    document.add_heading('Smerter', 2)

    # Best, verst, gjennomsnitt
    p = document.add_paragraph("Smerter siste uken (beste, gjennomsnitt, verste): " + data["plager-beste"] \
        + "/" + data["plager-gjsn"] + "/" + data["plager-verste"])
    
    # Varighet smerter
    if(data["plager-mer-enn-et-aar"] == "ja"):
        p.add_run("\nVarighet: " + data["plager-aar"] + " år")
    else:
        p.add_run("\nVarighet: " + data["plager-mnd"] + " måneder")

    # Fibroskjema
    document.add_heading('Fibromyalgiskjema', 4)
    p = document.add_paragraph("Utmattelse: " + data["fibro-utmattelse"].capitalize())
    write_var_snippet_and_var_code("fibro-kognisjon", data, codebook, p)
    write_var_snippet_and_var_code("fibro-trett", data, codebook, p)
    write_var_text_and_response("fibro-mage", data, codebook, p)
    write_var_text_and_response("fibro-depresjon", data, codebook, p)
    write_var_text_and_response("fibro-hodepine", data, codebook, p)
    write_var_text_report_and_multi_response_bullet("fibro-smerteomraader_1", data, codebook, document)

    # Tanker om smerter
    document.add_heading('Tanker om smertene', 4)
    document.add_paragraph("0 = 'helt uenig'. 10 = 'helt enig'")
    p = document.add_paragraph("", style = "List Bullet")
    p.add_run(codebook["skadelig"]["var_text"] + " (" + data["skadelig"] + ")").font.color.rgb = get_color_gradient(int(data["skadelig"]) / 10)
    p = document.add_paragraph("", style = "List Bullet")
    p.add_run(codebook["farlig"]["var_text"] + " (" + data["farlig"] + ")").font.color.rgb = get_color_gradient(int(data["farlig"]) / 10)
    p = document.add_paragraph("", style = "List Bullet")
    p.add_run(codebook["kroppen-skadet"]["var_text"] + " (" + data["kroppen-skadet"] + ")").font.color.rgb = get_color_gradient(int(data["kroppen-skadet"]) / 10)
    p = document.add_paragraph("", style = "List Bullet")
    p.add_run(codebook["unngaa-vondt"]["var_text"] + " (" + data["unngaa-vondt"] + ")").font.color.rgb = get_color_gradient(int(data["unngaa-vondt"]) / 10)
    p = document.add_paragraph("", style = "List Bullet")
    p.add_run(codebook["tro-paa-bedring"]["var_text"] + " (" + data["tro-paa-bedring"] + ")").font.color.rgb = get_color_gradient(1 - (int(data["tro-paa-bedring"]) / 10))

    # Tanker om smerter
    document.add_heading('Tanker og følelser når du opplever smerte', 6)
    p = document.add_paragraph("")
    write_var_text_and_response("fryktelig", data, codebook, p, newLine= False)
    write_var_text_and_response("forverring", data, codebook, p)
    write_var_text_and_response("ikke-ut-av-hodet", data, codebook, p)
    write_var_text_and_response("hjelpeloshet", data, codebook, p)

    p = document.add_paragraph("")
    p.add_run("Pasientens tanker om årsak til plagene: " + data["aarsak"])


def write_earlier_treatment(data, codebook, document):
    document.add_heading('Tidligere behandling', 4)
    p = document.add_paragraph("")
    write_var_text_report_and_multi_response("tidligere-behandling_1", data, codebook, p, colon=False, separator=", ", leading_newline=False)
    if len(data["annen-tidligere-beh"]) > 0:
        write_var_snippet_and_var_code("annen-tidligere-beh", data, codebook, p)

#endregion

##############
### ARBEID ###
##############
#region Arbeid

def write_work_related_t1v20(data, codebook, document):

    document.add_heading('Arbeidshistorikk og utdanning', 2)
    p = document.add_paragraph("")

    # Utdanning og stillingsforhold
    write_var_snippet_and_response("utdannelse", data, codebook, p, True, False)
    write_var_snippet_and_response("tid-i-arbeidslivet", data, codebook, p)
    write_var_snippet_and_response("arbeidsforhold", data, codebook, p)
    write_var_snippet_and_var_code("trivsel-jobb", data, codebook, p)    
    write_var_snippet_and_var_code("yrke-fritekst", data, codebook, p)    
    write_var_snippet_and_var_code("stillingsprosent", data, codebook, p)
    write_var_text_report_and_multi_response("sektor_1", data, codebook, p)

    # Sykemelding og AAP
    write_var_snippet_and_response("sm-aap", data, codebook, p, False)
    write_var_snippet_and_var_code("prosent-sykemeldt", data, codebook, p)
    write_var_snippet_and_var_code("fare-sykemelding", data, codebook, p)
    write_var_snippet_and_var_code("fare-mer-sykemeldt", data, codebook, p)
    write_var_snippet_and_response("oppfolgingsplan", data, codebook, p)
    write_var_snippet_and_var_code("avtalt-i-oppfolgingsplan", data, codebook, p)
    write_var_snippet_and_response("aktivitetsplan", data, codebook, p)
    write_var_snippet_and_var_code("avtalt-i-aktivitetsplan", data, codebook, p)
    write_var_snippet_and_var_code("aarsak-sm-app", data, codebook, p, True, True)
    write_var_snippet_and_response("samarbeid-nav", data, codebook, p)
    write_var_snippet_and_response("tidl-tiltak", data, codebook, p)
    write_var_snippet_and_response("rapport-avklaring", data, codebook, p)
    write_var_text_report_and_multi_response("ytelser_1", data, codebook, p)
    write_var_snippet_and_response("varighet-sm-siste-ar", data, codebook, p)

    write_var_snippet_and_response("okonomi", data, codebook, p)
    write_var_snippet_and_response("sokt-ufor", data, codebook, p)
    write_var_snippet_and_response("erstatningssak", data, codebook, p)
 
    # Vurdering av jobbmestring
    write_var_snippet_and_response("estimat-rtw", data, codebook, p)
    write_var_snippet_and_response("onsket-jobb", data, codebook, p)
    write_var_snippet_and_response("arbeidsevne-fysiske-krav", data, codebook, p)
    write_var_snippet_and_response("arbeidsevne-mentale-krav", data, codebook, p)
    write_var_snippet_and_response("arbeidsevne-sosiale-krav", data, codebook, p)
    write_var_snippet_and_var_code("arbeidsevne-ny-jobb", data, codebook, p)
    write_var_snippet_and_var_code("arbeidsevne-aktuell-jobb", data, codebook, p)
    write_var_snippet_and_var_code("endringer-jobbsit-rtw", data, codebook, p)
   


def write_work_related_t2v10(data, codebook, document):

    document.add_heading('Arbeid', 2)
    p = document.add_paragraph("")

    # Stillingsforhold
    write_var_snippet_and_response("arbeidsforhold", data, codebook, p, newLine=False)
    write_var_snippet_and_var_code("trivsel-jobb", data, codebook, p)    
    write_var_snippet_and_var_code("stillingsprosent", data, codebook, p)
    
    # Sykemelding og AAP
    write_var_snippet_and_response("sm-aap", data, codebook, p, False)
    write_var_snippet_and_var_code("prosent-sykemeldt", data, codebook, p)
    write_var_snippet_and_var_code("fare-sykemelding", data, codebook, p)
    write_var_snippet_and_var_code("fare-mer-sykemeldt", data, codebook, p)
    write_var_snippet_and_response("samarbeid-nav", data, codebook, p)
    
    write_var_snippet_and_response("okonomi", data, codebook, p)
    write_var_snippet_and_response("sokt-ufor", data, codebook, p)
    
    # Vurdering av jobbmestring
    write_var_snippet_and_response("estimat-rtw", data, codebook, p)
    write_var_snippet_and_response("onsket-jobb", data, codebook, p)
    write_var_snippet_and_response("arbeidsevne-fysiske-krav", data, codebook, p)
    write_var_snippet_and_response("arbeidsevne-mentale-krav", data, codebook, p)
    write_var_snippet_and_response("arbeidsevne-sosiale-krav", data, codebook, p)
    write_var_snippet_and_var_code("arbeidsevne-ny-jobb", data, codebook, p)
    write_var_snippet_and_var_code("arbeidsevne-aktuell-jobb", data, codebook, p)
   

#endregion


####################
### LIVSKVALITET ###
####################
def write_quality_of_life(data, codebook, document):
    document.add_heading('Livskvalitet', 2)
    p = document.add_paragraph("");
    write_var_snippet_and_var_code("livskvalitet", data, codebook, p, newLine=False)


###############################################
### FYSISK AKTIVITET, PSYKISK HELSE OG SØVN ###
###############################################
#region Fysisk aktivitet, psykisk helse og søvn

    # FYSISK AKTIVITET OG KOSTHOLD
def write_exercise_and_more(data, codebook, document):
    document.add_heading("Mosjon, BMI og kosthold", 2)
    p = document.add_paragraph("")
    vekt = int(data["vekt"])
    hoyde = float(data["hoyde"]) / 100
    bmi = vekt / (hoyde * hoyde)
    p.add_run(f"BMI: {bmi:.1f}")
    write_var_snippet_and_response("mosjon", data, codebook, p, True)
    write_var_snippet_and_response("kosthold", data, codebook, p)
    write_var_snippet_and_response("maaltidsrytme", data, codebook, p)
    if "vurderer-endre_1" in data.keys():
        write_var_text_report_and_multi_response("vurderer-endre_1", data, codebook, p)


    # SCL
def write_scl(data, codebook, document):
    document.add_heading('Symptom checklist 10', 2)

    # samle alle keys for scl
    keys = []
    for key in data:
        if key[0:3] == "scl":
            keys.append(key)


    ikke = ""
    litt = ""
    ganske = ""
    veldig = ""

   # Fordele ledd på avsnitt etter respons
    for key in keys:
        if data[key] == "veldig-mye":
            if(len(veldig) > 0):
                veldig = veldig + "\n"
            veldig = veldig + (codebook[key]["var_text"])
        elif data[key] == "ganske-mye":
            if(len(ganske) > 0):
                ganske = ganske + "\n"
            ganske = ganske + (codebook[key]["var_text"])
        elif data[key] == "litt-plaget":
            if(len(litt) > 0):
                litt = litt + "\n"
            litt = litt + (codebook[key]["var_text"])
        else:
            if(len(ikke) > 0):
                ikke = ikke + "\n"
            ikke = ikke + (codebook[key]["var_text"])

    if len(veldig) > 0:
            document.add_heading("Respons: Veldig mye", 3)
            document.add_paragraph(veldig)
    if len(ganske) > 0:
            document.add_heading("Respons: Ganske mye", 3)
            document.add_paragraph(ganske)
    if len(litt) > 0:
            document.add_heading("Respons: Litt plager", 3)
            document.add_paragraph(litt)
    if len(ikke) > 0:
            document.add_heading("Respons: Ikke plaget", 3)
            document.add_paragraph(ikke)
    


def write_isi(data, codebook, document):
    document.add_heading('Insomnia Severity Scale', 2)
    p = document.add_paragraph("")
    first = True
    for key in data:
        if key[0:3] == "isi":
            if first:
                first = False
                write_var_text_and_response(key, data, codebook, p, newLine=False)
            else:
                write_var_text_and_response(key, data, codebook, p)


#endregion


###############################
### HJELPERE FOR ENKELTLEDD ###
###############################


def write_var_text_and_raw_response(var, data, codebook, paragraph, colon=True, newLine = True):
    # skip if value is missing
    if(data[var] == ""):
        return

    s = "\n" if newLine else ""
    s += codebook[var]["var_text"]
    
    if(colon):
        s += ": "
    else:
        s += " "
    
    s += data[var]
    paragraph.add_run(s)


def write_response(var, data, codebook, document):
    s = codebook[var]["responses"][data[var]]
    document.add_paragraph(s + ".")


def write_var_snippet_and_var_code(var, data, codebook, paragraph, colon=True, capitalization = True, newLine = True):
    # skip if value is missing
    if(data[var] == ""):
        return

    s = "\n" if newLine else ""

    s += codebook[var]["var_text_report"]
    
    if(colon):
        s += ": "
    else:
        s += " "
    
    s += data[var].capitalize()

    paragraph.add_run(s)

def write_var_snippet_and_response(var, data, codebook, paragraph, colon=True, newLine = True):
    # skip if value is missing
    if(data[var] == ""):
        return

    s = "\n" if newLine else ""
    s += codebook[var]["var_text_report"]
    
    if(colon):
        s += ": "
    else:
        s += " "
    
    response_code = data[var]
    s += codebook[var]["responses"][response_code]

    paragraph.add_run(s)

def write_var_text_and_response(var, data, codebook, paragraph, colon=True, newLine = True):
    # skip if value is missing
    if(data[var] == ""):
        return

    s = "\n" if newLine else ""
    s += codebook[var]["var_text"]
    
    if(colon):
        s += ": "
    else:
        s += " "
    
    response_code = data[var]
    s += codebook[var]["responses"][response_code]

    paragraph.add_run(s)


def write_var_text_report_and_multi_response_bullet(var, data, codebook, document, colon=True):   
    s = codebook[var]["var_text_report"]
    if(colon):
        s += ": "
    elif len(codebook[var]["var_text_report"]) > 0:
        s += " "
    document.add_heading(s, 5)

    num_of_responses = len(codebook[var]["responses"])
    base_var = var[:-2]
    for i in range(1, num_of_responses + 1):
        respons_var = base_var + "_" + str(i)
        if(data[respons_var] != ""):
            document.add_paragraph(codebook[respons_var]["responses"][data[respons_var]], style='List Bullet')



def write_var_text_report_and_multi_response(var, data, codebook, p, colon=True, leading_newline = True, separator = "; "):   
    s = "\n" if leading_newline else ""
    s += codebook[var]["var_text_report"]
    if(colon):
        s += ": "
    elif len(codebook[var]["var_text_report"]) > 0:
        s += " "

    response = ""
    num_of_responses = len(codebook[var]["responses"])
    base_var = var[:-2]
    for i in range(1, num_of_responses + 1):
        respons_var = base_var + "_" + str(i)
        if(data[respons_var] != ""):
            if(len(response) == 0):
                response += codebook[respons_var]["responses"][data[respons_var]]
            else:
                response += separator + codebook[respons_var]["responses"][data[respons_var]].lower()

    s += response

    if(len(response) > 0):
        p.add_run(s)



###################
### COMPUTERING ###
###################
#region Computering

def scl_score(data):
    points = 0
    for key in data:
        if (key[0:3] == "scl"):
            val = data[key]
            if val == "ikke-plaget": points += 1
            elif val == "litt-plaget": points += 2
            elif val == "ganske-mye": points += 3
            elif val == "veldig-mye": points += 4
            else: print("error, scl response")
    return points / 10

def isi_score(data):
    points = 0
    for key in data:
        if (key[0:3] == "isi"):
            val = data[key]
            if val == "veldig-tilfreds" or val == "ingen": continue
            elif val == "milde" or val == "tilfreds" or val == "litt": points += 1
            elif val == "moderate" or val == "verken" or val == "noe": points += 2
            elif val == "alvorlige" or val == "utilfreds" or val == "ganske-mye" or val == "ganske" or val == "mye": points += 3
            elif val == "veldig-alvorlige" or val == "veldig-utilfreds" or val == "veldig" or val == "veldig-mye": points += 4
            else: print("error, isi response: " + val)
    return points

def wpi_score(data):
    counter = 0
    for key in data:
        if (key[:21] == "fibro-smerteomraader_"):
            if len(data[key]) > 0:
                counter += 1
    return counter

def sss_score(data, codebook):
    counter = 0
    counter += int(codebook["fibro-utmattelse"]["responses"][data["fibro-utmattelse"]])
    counter += int(codebook["fibro-utmattelse"]["responses"][data["fibro-kognisjon"]])
    counter += int(codebook["fibro-utmattelse"]["responses"][data["fibro-trett"]])
    counter += 1 if data["fibro-mage"] == "ja" else 0
    counter += 1 if data["fibro-depresjon"] == "ja" else 0
    counter += 1 if data["fibro-hodepine"] == "ja" else 0
    return counter

def number_of_pain_regions(data):
    # samle alle fibrokoder i en array
    smerteomraader = []
    for key in data:
        if (key[:21] == "fibro-smerteomraader_"):
            smerteomraader.append(data[key])

    counter = 0

    # Øvre venstre
    if "skulder-ven" in smerteomraader or "overarm-ven" in smerteomraader or "underarm-ven" in smerteomraader:
        counter += 1
    # Øvre høyre
    if "skulder-hoy" in smerteomraader or "overarm-hoy" in smerteomraader or "underarm-hoy" in smerteomraader:
        counter += 1
    # Nedre venstre
    if "hofte-ven" in smerteomraader or "legg-fot-ven" in smerteomraader or "laar-kne-ven" in smerteomraader:
        counter += 1
    # Nedre høyre
    if "hofte-hoy" in smerteomraader or "legg-fot-hoy" in smerteomraader or "laar-kne-hoy" in smerteomraader:
        counter += 1
    # Midtregion
    if "ovre-rygg" in smerteomraader or "korsrygg" in smerteomraader or "nakke-hals" in smerteomraader:
        counter += 1
    
    return counter

    

# val må være fra 0 - 1. 0 gir grønn og 1 gir rød
def get_color_gradient(val):
     if val <= .1:
         return RGBColor(0, 128, 0)
     elif val <= .2:
         return RGBColor(64, 128, 0)
     elif val <= .3:
         return RGBColor(128, 128, 0)
     elif val <= .4:
         return RGBColor(128, 128, 0)
     elif val <= .5:
         return RGBColor(128, 96, 0)
     elif val <= .6:
         return RGBColor(128, 96, 0)
     elif val <= .7:
         return RGBColor(128, 64, 0)
     elif val <= .8:
         return RGBColor(128, 64, 0)
     elif val <= .9:
         return RGBColor(128, 0, 0)
     else:
         return RGBColor(128, 0, 0)
    


def oppfyller_fibrokriterier(data, codebook):
    wpi_and_sss = (wpi_score(data) >= 7 and sss_score(data, codebook) >= 5) or (wpi_score(data) >= 4 and sss_score(data, codebook) >= 9)
    mer_enn_1_aar = data["plager-mer-enn-et-aar"] == "ja" or int(data["plager-mnd"]) >= 3 
    return wpi_and_sss and mer_enn_1_aar and (number_of_pain_regions(data) >= 4)

#endregion


def main():
        print("Main function ikke implementert")


if __name__ == "__main__":
    main()
