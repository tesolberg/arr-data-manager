from docx import Document
from docx.shared import Inches
import csv
import locale

def main():
    row =  get_first_row("data.csv")
    generate_report(row, cb)


# dict av enkeltbesvarelse -> void + rapport i word med filnavn = checksum
def generate_report(data, codebook):
    document = Document()

    document.add_heading('Rapport fra nettbasert kartlegging - Arbeidsrettet rehabilitering', 0)

    document.add_heading('Arbeidshistorikk og utdanning')

    for variable in codebook.data:
        if(variable == "fnr"):
            continue
        print(variable)
        var_report_format = codebook.data[variable]["data-name"]
        document.add_paragraph (var_report_format + ": " + codebook.data[variable]["responses"][data[variable]])
 



    document.save('demo.docx')



def get_first_row(path):
    with open(path, newline="") as csvfile:
        reader = csv.DictReader(csvfile, dialect="excel-tab")
        row = reader.__next__()
        return row


if __name__ == "__main__":
    main()
